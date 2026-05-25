"""
Whisper Studio — локальная обёртка для транскрипции аудио на базе faster-whisper.

Запуск:   python app.py
Откроется в браузере:  http://127.0.0.1:8756

Всё работает на вашем компьютере. Аудио никуда не отправляется.
"""

import os
import sys
from pathlib import Path as _Path

# Если запущены под pythonw.exe (нет консоли), stdout/stderr = None.
# Многие библиотеки падают при попытке write() — перенаправляем в лог.
if sys.stdout is None or sys.stderr is None:
    _log_dir = _Path(os.environ.get("LOCALAPPDATA", _Path.home())) / "WhisperStudio"
    try:
        _log_dir.mkdir(parents=True, exist_ok=True)
        _log_file = open(_log_dir / "app.log", "a", encoding="utf-8", buffering=1)
        sys.stdout = _log_file
        sys.stderr = _log_file
    except Exception:
        import io
        sys.stdout = io.StringIO()
        sys.stderr = sys.stdout

# Подключаем NVIDIA cuBLAS / cuDNN, установленные через pip (для GPU-режима).
if sys.platform == "win32":
    _site = _Path(sys.prefix) / "Lib" / "site-packages" / "nvidia"
    for _sub in ("cublas", "cudnn", "cuda_nvrtc"):
        _bin = _site / _sub / "bin"
        if _bin.is_dir():
            os.add_dll_directory(str(_bin))
            os.environ["PATH"] = str(_bin) + os.pathsep + os.environ.get("PATH", "")

import json
import uuid
import time
import threading
import datetime
import traceback
import webbrowser
from pathlib import Path

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse, Response, FileResponse
from fastapi.staticfiles import StaticFiles

# --------------------------------------------------------------------------
# Конфигурация
# --------------------------------------------------------------------------
BASE = Path(__file__).parent.resolve()


def _resolve_data_dir() -> Path:
    """Где хранить пользовательские данные.

    - Если задан env WHISPER_STUDIO_USER_DATA — используем его.
    - Если BASE лежит внутри Program Files — пишем в %APPDATA%\Whisper Studio.
    - Иначе (dev-режим) — рядом с исходниками: BASE/data.
    """
    if os.environ.get("WHISPER_STUDIO_USER_DATA"):
        return Path(os.environ["WHISPER_STUDIO_USER_DATA"]).resolve()
    pf_paths = []
    for var in ("PROGRAMFILES", "PROGRAMFILES(X86)", "PROGRAMW6432"):
        v = os.environ.get(var)
        if v:
            pf_paths.append(Path(v).resolve())
    base_str = str(BASE).lower()
    if any(base_str.startswith(str(p).lower()) for p in pf_paths):
        return Path(os.environ.get("APPDATA", str(Path.home()))) / "Whisper Studio"
    return BASE / "data"


DATA = _resolve_data_dir()
UPLOADS = DATA / "uploads"
RECORDINGS = DATA / "recordings"
for _d in (UPLOADS, RECORDINGS):
    _d.mkdir(parents=True, exist_ok=True)

HOST = "127.0.0.1"
PORT = 8756
DEFAULT_MODEL = "medium"          # small | medium | large-v3
AVAILABLE_MODELS = ["small", "medium", "large-v3"]
APP_VERSION = "1.5"
GITHUB_REPO = "urtenovcom/whisper-studio"

app = FastAPI(title="Whisper Studio")

# --------------------------------------------------------------------------
# Определение устройства (GPU/CPU) и кэш модели
# --------------------------------------------------------------------------
def detect_device():
    """Возвращает (device, compute_type). Пытается использовать GPU NVIDIA."""
    try:
        import ctranslate2
        if ctranslate2.get_cuda_device_count() > 0:
            return "cuda", "float16"
    except Exception:
        pass
    return "cpu", "int8"


DEVICE, COMPUTE_TYPE = detect_device()

_model_cache = {}
_model_lock = threading.Lock()


def get_model(size):
    """Загружает модель faster-whisper (с кэшем). Держим в памяти одну модель."""
    with _model_lock:
        if size not in _model_cache:
            from faster_whisper import WhisperModel
            _model_cache.clear()  # экономим память: одна модель за раз
            _model_cache[size] = WhisperModel(
                size, device=DEVICE, compute_type=COMPUTE_TYPE
            )
        return _model_cache[size]


# --------------------------------------------------------------------------
# Диаризация — кто говорит (sherpa-onnx, без HuggingFace и без токенов)
# --------------------------------------------------------------------------
# Модели диаризации идут в комплекте, лежат рядом с приложением (read-only ок).
_DIAR_CANDIDATES = [BASE / "data" / "models" / "diarize", DATA / "models" / "diarize"]
DIAR_DIR = next((p for p in _DIAR_CANDIDATES if p.exists()), _DIAR_CANDIDATES[0])
DIAR_SEG = DIAR_DIR / "segmentation.onnx"
DIAR_EMB = DIAR_DIR / "embedding_v2.onnx"
if not DIAR_EMB.exists():
    DIAR_EMB = DIAR_DIR / "embedding.onnx"


def diarization_available():
    try:
        import sherpa_onnx  # noqa: F401
        return DIAR_SEG.exists() and DIAR_EMB.exists()
    except Exception:
        return False


def _load_audio_16k_mono(audio_path):
    """Декодирует любой аудио/видео-файл в float32 моно 16 кГц."""
    import av
    import numpy as np

    container = av.open(str(audio_path))
    stream = container.streams.audio[0]
    resampler = av.AudioResampler(format="flt", layout="mono", rate=16000)
    samples = []
    for frame in container.decode(stream):
        for new_frame in resampler.resample(frame):
            samples.append(new_frame.to_ndarray().flatten())
    for new_frame in resampler.resample(None):
        samples.append(new_frame.to_ndarray().flatten())
    container.close()
    if not samples:
        return np.zeros(0, dtype=np.float32)
    return np.concatenate(samples).astype(np.float32)


def apply_diarization(audio_path, segments, num_speakers=0):
    """Размечает сегменты по говорящим через sherpa-onnx (локально, без сети).

    num_speakers: 0 = авто-кластеризация по порогу; >0 = ровно столько групп.
    """
    import sherpa_onnx

    if num_speakers and num_speakers > 0:
        clustering = sherpa_onnx.FastClusteringConfig(
            num_clusters=int(num_speakers), threshold=0.5
        )
    else:
        clustering = sherpa_onnx.FastClusteringConfig(num_clusters=-1, threshold=0.65)

    config = sherpa_onnx.OfflineSpeakerDiarizationConfig(
        segmentation=sherpa_onnx.OfflineSpeakerSegmentationModelConfig(
            pyannote=sherpa_onnx.OfflineSpeakerSegmentationPyannoteModelConfig(
                model=str(DIAR_SEG),
            ),
        ),
        embedding=sherpa_onnx.SpeakerEmbeddingExtractorConfig(model=str(DIAR_EMB)),
        clustering=clustering,
        min_duration_on=0.5,
        min_duration_off=0.5,
    )
    sd = sherpa_onnx.OfflineSpeakerDiarization(config)
    samples = _load_audio_16k_mono(audio_path)
    result = sd.process(samples).sort_by_start_time()
    turns = [(r.start, r.end, r.speaker) for r in result]

    labelmap = {}
    for seg in segments:
        mid = (seg["start"] + seg["end"]) / 2.0
        chosen = None
        for st, en, label in turns:
            if st <= mid <= en:
                chosen = label
                break
        if chosen is None and turns:
            chosen = min(
                turns, key=lambda t: min(abs(t[0] - mid), abs(t[1] - mid))
            )[2]
        if chosen is not None:
            if chosen not in labelmap:
                labelmap[chosen] = f"Speaker {len(labelmap) + 1}"
            seg["speaker"] = labelmap[chosen]
    return segments, len(labelmap)


# --------------------------------------------------------------------------
# Фоновые задачи транскрипции
# --------------------------------------------------------------------------
jobs = {}  # job_id -> dict
jobs_lock = threading.Lock()


def _set_job(job_id, **kw):
    with jobs_lock:
        jobs.setdefault(job_id, {}).update(kw)


def run_transcription(job_id, audio_path, display_name, language, model_size, diarize, num_speakers=0, project=None):
    try:
        _set_job(job_id, status="loading", message="Загрузка модели…", progress=0.0)
        model = get_model(model_size)

        _set_job(job_id, status="transcribing", message="Распознавание речи…")
        lang = None if language in ("auto", "", None) else language
        segments_gen, info = model.transcribe(
            str(audio_path), language=lang, vad_filter=True, beam_size=5
        )
        duration = float(info.duration or 0)

        segments = []
        for s in segments_gen:
            with jobs_lock:
                if jobs.get(job_id, {}).get("cancelled"):
                    break
            segments.append(
                {
                    "start": float(s.start),
                    "end": float(s.end),
                    "text": s.text.strip(),
                    "speaker": None,
                }
            )
            if duration > 0:
                _set_job(job_id, progress=min(0.98, s.end / duration))

        with jobs_lock:
            if jobs.get(job_id, {}).get("cancelled"):
                _set_job(job_id, status="cancelled", message="Отменено", progress=0)
                try:
                    audio_path.unlink(missing_ok=True)
                except Exception:
                    pass
                return

        speakers = 0
        diar_warning = None
        if diarize:
            _set_job(job_id, status="diarizing", message="Определение говорящих…")
            try:
                segments, speakers = apply_diarization(audio_path, segments, num_speakers)
            except Exception as e:
                diar_warning = f"Диаризация пропущена: {e}"

        rec_id = save_recording(
            display_name, audio_path, segments, info, model_size, speakers, project
        )
        _set_job(
            job_id,
            status="done",
            progress=1.0,
            message="Готово",
            recording_id=rec_id,
            diar_warning=diar_warning,
        )
    except Exception as e:
        _set_job(
            job_id,
            status="error",
            error=str(e),
            trace=traceback.format_exc(),
        )


PROJECTS_FILE = DATA / "projects.json"


def load_projects():
    if not PROJECTS_FILE.exists():
        return []
    try:
        return json.loads(PROJECTS_FILE.read_text(encoding="utf-8"))
    except Exception:
        return []


def save_projects(projects):
    PROJECTS_FILE.write_text(
        json.dumps(projects, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def save_recording(name, audio_path, segments, info, model_size, speakers, project=None):
    rec_id = uuid.uuid4().hex[:12]
    rec = {
        "id": rec_id,
        "name": name,
        "created": datetime.datetime.now().isoformat(timespec="seconds"),
        "duration": float(info.duration or 0),
        "language": info.language,
        "model": model_size,
        "speakers": speakers,
        "audio": audio_path.name,
        "project": project,
        "segments": segments,
    }
    (RECORDINGS / f"{rec_id}.json").write_text(
        json.dumps(rec, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return rec_id


def load_recording(rec_id):
    path = RECORDINGS / f"{rec_id}.json"
    if not path.exists():
        raise HTTPException(404, "Запись не найдена")
    return json.loads(path.read_text(encoding="utf-8"))


# --------------------------------------------------------------------------
# Экспорт
# --------------------------------------------------------------------------
def fmt_ts(t, sep=","):
    t = max(0.0, float(t))
    h = int(t // 3600)
    m = int((t % 3600) // 60)
    s = int(t % 60)
    ms = int(round((t - int(t)) * 1000))
    return f"{h:02d}:{m:02d}:{s:02d}{sep}{ms:03d}"


def export_srt(rec):
    out = []
    for i, s in enumerate(rec["segments"], 1):
        spk = f"[{s['speaker']}] " if s.get("speaker") else ""
        out.append(
            f"{i}\n{fmt_ts(s['start'])} --> {fmt_ts(s['end'])}\n{spk}{s['text']}\n"
        )
    return "\n".join(out)


def export_vtt(rec):
    out = ["WEBVTT", ""]
    for s in rec["segments"]:
        spk = f"<v {s['speaker']}>" if s.get("speaker") else ""
        out.append(f"{fmt_ts(s['start'], '.')} --> {fmt_ts(s['end'], '.')}")
        out.append(f"{spk}{s['text']}")
        out.append("")
    return "\n".join(out)


def export_txt(rec):
    lines = []
    last_speaker = None
    for s in rec["segments"]:
        spk = s.get("speaker")
        if spk and spk != last_speaker:
            lines.append(f"\n{spk}:")
            last_speaker = spk
        lines.append(s["text"])
    return "\n".join(lines).strip() + "\n"


def export_docx(rec):
    import io
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(rec["name"], level=1)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Дата: {rec['created']}   Длительность: {fmt_ts(rec['duration'], '.')}   "
        f"Язык: {rec.get('language', '-')}   Модель: {rec.get('model', '-')}"
    ).italic = True

    last_speaker = None
    for s in rec["segments"]:
        spk = s.get("speaker")
        if spk and spk != last_speaker:
            p = doc.add_paragraph()
            p.add_run(spk).bold = True
            last_speaker = spk
        p = doc.add_paragraph(s["text"])
        p.paragraph_format.space_after = Pt(4)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------
# API
# --------------------------------------------------------------------------
@app.get("/api/about")
def api_about():
    return {"version": APP_VERSION, "repo": GITHUB_REPO}


# --------------------------------------------------------------------------
# Автозапуск при старте Windows (реестр HKCU\...Run, без прав администратора)
# --------------------------------------------------------------------------
_AUTOSTART_KEY = r"Software\Microsoft\Windows\CurrentVersion\Run"
_AUTOSTART_NAME = "WhisperStudio"


def _autostart_command() -> str:
    pythonw = _Path(sys.executable).with_name("pythonw.exe")
    if not pythonw.exists():
        pythonw = _Path(sys.executable)
    return f'"{pythonw}" "{BASE / "app.py"}" --minimized'


def autostart_is_enabled() -> bool:
    if sys.platform != "win32":
        return False
    try:
        import winreg
        with winreg.OpenKey(winreg.HKEY_CURRENT_USER, _AUTOSTART_KEY) as k:
            val, _ = winreg.QueryValueEx(k, _AUTOSTART_NAME)
            return bool(val)
    except FileNotFoundError:
        return False
    except Exception:
        return False


def autostart_set(enabled: bool):
    if sys.platform != "win32":
        return
    import winreg
    with winreg.OpenKey(
        winreg.HKEY_CURRENT_USER, _AUTOSTART_KEY, 0, winreg.KEY_SET_VALUE
    ) as k:
        if enabled:
            winreg.SetValueEx(k, _AUTOSTART_NAME, 0, winreg.REG_SZ, _autostart_command())
        else:
            try:
                winreg.DeleteValue(k, _AUTOSTART_NAME)
            except FileNotFoundError:
                pass


@app.get("/api/autostart")
def api_autostart_get():
    return {"enabled": autostart_is_enabled()}


@app.put("/api/autostart")
async def api_autostart_set(payload: dict):
    enabled = bool((payload or {}).get("enabled"))
    autostart_set(enabled)
    return {"enabled": autostart_is_enabled()}


def _version_tuple(v: str):
    parts = []
    for p in (v or "").lstrip("vV").split("."):
        try:
            parts.append(int(p))
        except Exception:
            parts.append(0)
    return tuple(parts)


@app.get("/api/update/check")
def api_update_check():
    try:
        import requests
        r = requests.get(
            f"https://api.github.com/repos/{GITHUB_REPO}/releases/latest",
            timeout=8,
            headers={"Accept": "application/vnd.github+json"},
        )
        if r.status_code != 200:
            return {"available": False, "error": f"HTTP {r.status_code}"}
        data = r.json()
        latest_tag = (data.get("tag_name") or "").lstrip("vV")
        if not latest_tag:
            return {"available": False, "error": "no version"}
        try:
            available = _version_tuple(latest_tag) > _version_tuple(APP_VERSION)
        except Exception:
            available = latest_tag != APP_VERSION
        installer_url = None
        for asset in data.get("assets") or []:
            name = (asset.get("name") or "").lower()
            if name.endswith(".exe") and "setup" in name:
                installer_url = asset.get("browser_download_url")
                break
        return {
            "available": available,
            "version": latest_tag,
            "current": APP_VERSION,
            "download_url": installer_url,
            "notes": (data.get("body") or "")[:1500],
        }
    except Exception as e:
        return {"available": False, "error": str(e)}


@app.post("/api/update/install")
def api_update_install(payload: dict):
    url = (payload or {}).get("url")
    if not url:
        raise HTTPException(400, "url обязателен")
    import requests, tempfile, subprocess, threading, time as _time
    tmp = _Path(tempfile.gettempdir()) / "WhisperStudio-Update.exe"
    try:
        with requests.get(url, stream=True, timeout=60) as resp:
            resp.raise_for_status()
            with open(tmp, "wb") as f:
                for chunk in resp.iter_content(chunk_size=1024 * 1024):
                    if chunk:
                        f.write(chunk)
    except Exception as e:
        raise HTTPException(500, f"Не удалось скачать: {e}")
    pythonw = (_Path(sys.executable).with_name("pythonw.exe"))
    if not pythonw.exists():
        pythonw = _Path(sys.executable)
    helper = _Path(tempfile.gettempdir()) / "ws-update.bat"
    helper.write_text(
        '@echo off\n'
        'timeout /t 2 /nobreak >nul\n'
        f'"{tmp}" /VERYSILENT /SUPPRESSMSGBOXES /NORESTART\n'
        f'start "" "{pythonw}" "{BASE / "app.py"}"\n',
        encoding="utf-8",
    )
    DETACHED = 0x00000008
    subprocess.Popen(
        ["cmd.exe", "/c", str(helper)],
        creationflags=DETACHED, close_fds=True,
    )
    def _quit():
        _time.sleep(0.4)
        os._exit(0)
    threading.Thread(target=_quit, daemon=True).start()
    return {"ok": True}


@app.get("/api/status")
def api_status():
    return {
        "device": DEVICE,
        "compute_type": COMPUTE_TYPE,
        "models": AVAILABLE_MODELS,
        "default_model": DEFAULT_MODEL,
        "diarization": diarization_available(),
    }


@app.post("/api/upload")
async def api_upload(file: UploadFile = File(...)):
    file_id = uuid.uuid4().hex[:12]
    ext = Path(file.filename or "audio").suffix or ".bin"
    safe_name = f"{file_id}{ext}"
    dest = UPLOADS / safe_name
    with open(dest, "wb") as f:
        while chunk := await file.read(1024 * 1024):
            f.write(chunk)
    return {"file": safe_name, "name": file.filename or safe_name}


@app.post("/api/transcribe")
def api_transcribe(
    file: str = Form(...),
    name: str = Form(...),
    language: str = Form("auto"),
    model: str = Form(DEFAULT_MODEL),
    diarize: bool = Form(False),
    num_speakers: int = Form(0),
    project: str = Form(""),
):
    audio_path = UPLOADS / file
    if not audio_path.exists():
        raise HTTPException(404, "Файл не найден")
    if model not in AVAILABLE_MODELS:
        model = DEFAULT_MODEL

    print(
        f"[transcribe] file={file} model={model} lang={language} "
        f"diarize={diarize} num_speakers={num_speakers}",
        flush=True,
    )
    job_id = uuid.uuid4().hex[:12]
    _set_job(job_id, status="queued", progress=0.0, message="В очереди…")
    threading.Thread(
        target=run_transcription,
        args=(job_id, audio_path, name, language, model, diarize, num_speakers, project or None),
        daemon=True,
    ).start()
    return {"job": job_id}


@app.get("/api/job/{job_id}")
def api_job(job_id: str):
    with jobs_lock:
        job = jobs.get(job_id)
    if not job:
        raise HTTPException(404, "Задача не найдена")
    return job


@app.post("/api/job/{job_id}/cancel")
def api_cancel_job(job_id: str):
    with jobs_lock:
        if job_id in jobs:
            jobs[job_id]["cancelled"] = True
            return {"ok": True}
    raise HTTPException(404, "Задача не найдена")


@app.get("/api/recordings")
def api_recordings():
    items = []
    for path in RECORDINGS.glob("*.json"):
        try:
            rec = json.loads(path.read_text(encoding="utf-8"))
            speaker_set = []
            for s in rec.get("segments") or []:
                spk = s.get("speaker")
                if spk and spk not in speaker_set:
                    speaker_set.append(spk)
            items.append(
                {
                    "id": rec["id"],
                    "name": rec["name"],
                    "created": rec["created"],
                    "duration": rec["duration"],
                    "language": rec.get("language"),
                    "model": rec.get("model"),
                    "speakers": rec.get("speakers", 0),
                    "speaker_names": speaker_set,
                    "project": rec.get("project") or None,
                    "preview": (rec["segments"][0]["text"] if rec["segments"] else ""),
                }
            )
        except Exception:
            continue
    items.sort(key=lambda x: x["created"], reverse=True)
    return items


@app.get("/api/projects")
def api_projects():
    projects = load_projects()
    counts = {p: 0 for p in projects}
    no_project = 0
    for path in RECORDINGS.glob("*.json"):
        try:
            rec = json.loads(path.read_text(encoding="utf-8"))
            p = rec.get("project")
            if p and p in counts:
                counts[p] += 1
            else:
                no_project += 1
        except Exception:
            continue
    return {
        "projects": [{"name": p, "count": counts[p]} for p in projects],
        "no_project_count": no_project,
    }


@app.post("/api/projects")
async def api_create_project(payload: dict):
    name = (payload.get("name") or "").strip()
    if not name:
        raise HTTPException(400, "Имя проекта не может быть пустым")
    projects = load_projects()
    if name in projects:
        raise HTTPException(400, "Проект с таким именем уже существует")
    projects.append(name)
    save_projects(projects)
    return {"ok": True}


@app.put("/api/projects/{old_name}")
async def api_rename_project(old_name: str, payload: dict):
    new_name = (payload.get("name") or "").strip()
    if not new_name:
        raise HTTPException(400, "Имя не может быть пустым")
    projects = load_projects()
    if old_name not in projects:
        raise HTTPException(404, "Проект не найден")
    if new_name != old_name and new_name in projects:
        raise HTTPException(400, "Уже существует")
    idx = projects.index(old_name)
    projects[idx] = new_name
    save_projects(projects)
    for path in RECORDINGS.glob("*.json"):
        try:
            rec = json.loads(path.read_text(encoding="utf-8"))
            if rec.get("project") == old_name:
                rec["project"] = new_name
                path.write_text(
                    json.dumps(rec, ensure_ascii=False, indent=2), encoding="utf-8"
                )
        except Exception:
            continue
    return {"ok": True}


# --------------------------------------------------------------------------
# Диктовка (push-to-talk)
# --------------------------------------------------------------------------
import dictation as _dictation
_dictation.init(DATA / "dictation.json", get_model)


@app.get("/api/dictation/status")
def api_dictation_status():
    return _dictation.get_status()


@app.put("/api/dictation/settings")
async def api_dictation_settings(payload: dict):
    return _dictation.save_settings(payload or {})


@app.get("/api/dictation/microphones")
def api_dictation_microphones():
    return _dictation.list_microphones()


@app.delete("/api/projects/{name}")
def api_delete_project(name: str):
    projects = load_projects()
    if name not in projects:
        raise HTTPException(404, "Проект не найден")
    projects.remove(name)
    save_projects(projects)
    for path in RECORDINGS.glob("*.json"):
        try:
            rec = json.loads(path.read_text(encoding="utf-8"))
            if rec.get("project") == name:
                rec["project"] = None
                path.write_text(
                    json.dumps(rec, ensure_ascii=False, indent=2), encoding="utf-8"
                )
        except Exception:
            continue
    return {"ok": True}


@app.get("/api/recordings/{rec_id}")
def api_recording(rec_id: str):
    return load_recording(rec_id)


@app.put("/api/recordings/{rec_id}")
async def api_update_recording(rec_id: str, payload: dict):
    rec = load_recording(rec_id)
    if "name" in payload:
        rec["name"] = payload["name"]
    if "segments" in payload:
        rec["segments"] = payload["segments"]
        rec["speakers"] = len(
            {s["speaker"] for s in payload["segments"] if s.get("speaker")}
        )
    if "project" in payload:
        rec["project"] = (payload["project"] or None)
    (RECORDINGS / f"{rec_id}.json").write_text(
        json.dumps(rec, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return {"ok": True}


@app.delete("/api/recordings/{rec_id}")
def api_delete_recording(rec_id: str):
    rec = load_recording(rec_id)
    (RECORDINGS / f"{rec_id}.json").unlink(missing_ok=True)
    audio = UPLOADS / rec.get("audio", "")
    if audio.exists():
        audio.unlink(missing_ok=True)
    return {"ok": True}


@app.get("/api/recordings/{rec_id}/export")
def api_export(rec_id: str, format: str = "txt"):
    from urllib.parse import quote
    rec = load_recording(rec_id)
    base = "".join(c for c in rec["name"] if c.isalnum() or c in " _-").strip() or rec_id

    def _disp(ext):
        ascii_fallback = base.encode("ascii", "ignore").decode("ascii").strip() or rec_id
        return (
            f'attachment; filename="{ascii_fallback}.{ext}"; '
            f"filename*=UTF-8''{quote(base)}.{ext}"
        )

    if format == "srt":
        return Response(
            export_srt(rec),
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": _disp("srt")},
        )
    if format == "vtt":
        return Response(
            export_vtt(rec),
            media_type="text/vtt; charset=utf-8",
            headers={"Content-Disposition": _disp("vtt")},
        )
    if format == "docx":
        return Response(
            export_docx(rec),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": _disp("docx")},
        )
    # txt по умолчанию
    return Response(
        export_txt(rec),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": _disp("txt")},
    )


@app.get("/api/audio/{name}")
def api_audio(name: str):
    path = UPLOADS / name
    if not path.exists():
        raise HTTPException(404, "Аудио не найдено")
    return FileResponse(path)


# --------------------------------------------------------------------------
# Статика / фронтенд
# --------------------------------------------------------------------------
@app.get("/", response_class=HTMLResponse)
def index():
    return (BASE / "static" / "index.html").read_text(encoding="utf-8")


app.mount("/static", StaticFiles(directory=BASE / "static"), name="static")


# --------------------------------------------------------------------------
# Запуск
# --------------------------------------------------------------------------
if __name__ == "__main__":
    import uvicorn

    print("=" * 56)
    print("  Whisper Studio")
    print(f"  Устройство: {DEVICE} ({COMPUTE_TYPE})")
    print(f"  Откройте в браузере:  http://{HOST}:{PORT}")
    print("=" * 56)

    qt_app = None
    main_window = None
    tray = None
    try:
        import overlay as _overlay_mod
        if _overlay_mod.has_qt():
            from PySide6.QtWidgets import QApplication
            qt_app = QApplication.instance() or QApplication([])
            qt_app.setQuitOnLastWindowClosed(False)
            try:
                from mainwindow import apply_app_icon
                apply_app_icon(qt_app)
            except Exception as e:
                print(f"[main] icon apply failed: {e}", flush=True)
            _overlay_mod.init_on_main_thread()
    except Exception as e:
        print(f"[main] Qt unavailable: {e}", flush=True)

    if qt_app is not None:
        # uvicorn в фоновом потоке, Qt event loop на главном
        config = uvicorn.Config(app, host=HOST, port=PORT, log_level="warning")
        server = uvicorn.Server(config)
        srv_thread = threading.Thread(target=server.run, daemon=True)
        srv_thread.start()
        # Ждём, пока сервер начнёт принимать соединения, и открываем окно
        try:
            from mainwindow import MainWindow, setup_tray
            # Подождём пару секунд, чтобы uvicorn успел подняться
            from PySide6.QtCore import QTimer

            start_minimized = "--minimized" in sys.argv[1:]

            def _launch_window():
                global main_window, tray
                main_window = MainWindow(f"http://{HOST}:{PORT}/")
                tray = setup_tray(qt_app, main_window)
                if start_minimized:
                    # При автозапуске прячем — пользователь увидит только иконку в трее
                    main_window.hide()
                else:
                    main_window.show()
                    main_window.raise_()
                    main_window.activateWindow()
                    from PySide6.QtCore import QTimer as _QT
                    _QT.singleShot(1200, lambda: (
                        main_window.showNormal(),
                        main_window.raise_(),
                        main_window.activateWindow(),
                    ))

            QTimer.singleShot(800, _launch_window)
        except Exception as e:
            print(f"[main] window init error: {e}", flush=True)
        try:
            qt_app.exec()
        except KeyboardInterrupt:
            pass
    else:
        # Фолбэк: открываем в браузере
        def _open():
            time.sleep(1.5)
            try:
                webbrowser.open(f"http://{HOST}:{PORT}")
            except Exception:
                pass
        threading.Thread(target=_open, daemon=True).start()
        uvicorn.run(app, host=HOST, port=PORT, log_level="warning")
